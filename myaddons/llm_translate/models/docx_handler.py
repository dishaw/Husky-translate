"""Word document (.doc/.docx) parsing and reconstruction utilities.

Uses python-docx for .docx files.
For legacy .doc files, uses LibreOffice headless to convert to .docx first,
then processes the converted .docx with python-docx for full style fidelity.
"""

import base64
import io
import json
import logging
import math
import os
import re
import subprocess
import tempfile

_logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    _logger.warning("python-docx not installed. DOCX support unavailable.")


def estimate_tokens(text):
    """Estimate token count for a text string.

    Rough estimation: ~1 token per 4 chars for English, ~1 token per 1.5 chars for CJK.
    """
    if not text:
        return 0
    cjk_count = sum(1 for c in text if '\u4e00' <= c <= '\u9fff'
                    or '\u3400' <= c <= '\u4dbf'
                    or '\uf900' <= c <= '\ufaff'
                    or '\u3000' <= c <= '\u303f'
                    or '\u3040' <= c <= '\u309f'
                    or '\u30a0' <= c <= '\u30ff'
                    or '\uac00' <= c <= '\ud7af')
    ascii_count = len(text) - cjk_count
    return int(cjk_count / 1.5 + ascii_count / 4)


def split_long_paragraph(text, max_tokens=2000):
    """Split a paragraph that exceeds max_tokens into sentence-level chunks.

    Splits by common sentence terminators (。.！!？?；;) and regroups
    into sub-paragraphs each under max_tokens.

    Args:
        text: The paragraph text to potentially split.
        max_tokens: Maximum token threshold per chunk.

    Returns:
        list[str]: List of text segments, each under max_tokens.
    """
    if estimate_tokens(text) <= max_tokens:
        return [text]

    # Split by sentence terminators, keeping the delimiter
    sentence_pattern = r'(?<=[。.！!？?；;\n])'
    sentences = re.split(sentence_pattern, text)
    sentences = [s for s in sentences if s.strip()]

    if not sentences:
        return [text]

    chunks = []
    current_chunk = ""

    for sentence in sentences:
        test_chunk = current_chunk + sentence
        if estimate_tokens(test_chunk) > max_tokens and current_chunk:
            chunks.append(current_chunk)
            current_chunk = sentence
        else:
            current_chunk = test_chunk

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


# =========================================================================
# Word Auto-Numbering Engine
# =========================================================================

def _to_roman(num):
    """Convert integer to Roman numeral string."""
    vals = [
        (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
        (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
        (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I'),
    ]
    result = ''
    for v, s in vals:
        while num >= v:
            result += s
            num -= v
    return result


def _format_number(value, num_fmt):
    """Format a counter value according to the Word numbering format type."""
    if num_fmt in ('decimal', 'decimalZero'):
        return str(value)
    elif num_fmt == 'lowerLetter':
        return chr(ord('a') + (value - 1) % 26) if value >= 1 else str(value)
    elif num_fmt == 'upperLetter':
        return chr(ord('A') + (value - 1) % 26) if value >= 1 else str(value)
    elif num_fmt == 'lowerRoman':
        return _to_roman(value).lower()
    elif num_fmt == 'upperRoman':
        return _to_roman(value)
    elif num_fmt in ('chineseCounting', 'ideographTraditional',
                     'chineseCountingThousand', 'japaneseCounting'):
        cn = '零一二三四五六七八九十'
        if 1 <= value <= 10:
            return cn[value]
        elif 11 <= value <= 99:
            t, o = divmod(value, 10)
            return (cn[t] if t > 1 else '') + '十' + (cn[o] if o else '')
        return str(value)
    elif num_fmt == 'bullet':
        return '•'
    elif num_fmt == 'none':
        return ''
    return str(value)


def _build_numbering_engine(doc):
    """Build a numbering resolver from the document's numbering.xml definitions.

    Word auto-numbering (1., 2., a), i., 一、, etc.) is stored in XML as
    <w:numPr> within <w:pPr>, NOT as visible text in paragraph content.
    This engine parses numbering.xml to understand list formats, tracks
    counters sequentially, and resolves the visible numbering prefix for
    each paragraph.

    Args:
        doc: python-docx Document object.

    Returns:
        callable: resolve(para_element, para_style=None) -> (prefix_str, ilvl)
                  Returns (None, None) if paragraph has no numbering.
    """
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # ── Parse numbering definitions ──────────────────────────────
    abstract_nums = {}    # abstractNumId -> {ilvl -> {start, numFmt, lvlText}}
    num_to_abstract = {}  # numId -> abstractNumId
    num_overrides = {}    # numId -> {ilvl -> startOverride}

    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return lambda para_el, para_style=None: (None, None)
        numbering_xml = numbering_part.element
    except Exception:
        return lambda para_el, para_style=None: (None, None)

    for abstract in numbering_xml.findall(f'{{{w_ns}}}abstractNum'):
        abstract_id = abstract.get(f'{{{w_ns}}}abstractNumId')
        if abstract_id is None:
            continue
        levels = {}
        for lvl in abstract.findall(f'{{{w_ns}}}lvl'):
            ilvl_str = lvl.get(f'{{{w_ns}}}ilvl')
            if ilvl_str is None:
                continue
            start_el = lvl.find(f'{{{w_ns}}}start')
            fmt_el = lvl.find(f'{{{w_ns}}}numFmt')
            text_el = lvl.find(f'{{{w_ns}}}lvlText')
            levels[int(ilvl_str)] = {
                'start': int(start_el.get(f'{{{w_ns}}}val', '1'))
                         if start_el is not None else 1,
                'numFmt': fmt_el.get(f'{{{w_ns}}}val', 'decimal')
                          if fmt_el is not None else 'decimal',
                'lvlText': text_el.get(f'{{{w_ns}}}val', '%1.')
                           if text_el is not None else '%1.',
            }
        abstract_nums[abstract_id] = levels

    for num_el in numbering_xml.findall(f'{{{w_ns}}}num'):
        num_id = num_el.get(f'{{{w_ns}}}numId')
        abstract_ref = num_el.find(f'{{{w_ns}}}abstractNumId')
        if num_id and abstract_ref is not None:
            num_to_abstract[num_id] = abstract_ref.get(f'{{{w_ns}}}val')
        # Level overrides (restart numbering)
        for override in num_el.findall(f'{{{w_ns}}}lvlOverride'):
            olvl = override.get(f'{{{w_ns}}}ilvl')
            start_override = override.find(f'{{{w_ns}}}startOverride')
            if olvl is not None and start_override is not None:
                val = start_override.get(f'{{{w_ns}}}val')
                if val:
                    num_overrides.setdefault(num_id, {})[int(olvl)] = int(val)

    # ── Counter state (mutable, tracked across calls) ────────────
    counters = {}            # numId -> {ilvl -> current_value}
    last_ilvl_for_num = {}   # numId -> last ilvl seen

    def resolve(para_element, para_style=None):
        """Resolve the numbering prefix for a paragraph element.

        Args:
            para_element: lxml element (CT_P) of the paragraph.
            para_style: python-docx Style object (optional, for style-inherited numbering).

        Returns:
            tuple: (prefix_string, indent_level) or (None, None).
        """
        # Look for <w:numPr> in paragraph's own <w:pPr>
        pPr = para_element.find(f'{{{w_ns}}}pPr')
        numPr = pPr.find(f'{{{w_ns}}}numPr') if pPr is not None else None

        # If not in paragraph, check style chain
        if numPr is None and para_style is not None:
            style = para_style
            while style is not None:
                try:
                    style_pPr = style.element.find(f'{{{w_ns}}}pPr')
                    if style_pPr is not None:
                        numPr = style_pPr.find(f'{{{w_ns}}}numPr')
                        if numPr is not None:
                            break
                    style = style.base_style
                except Exception:
                    break

        if numPr is None:
            return (None, None)

        ilvl_el = numPr.find(f'{{{w_ns}}}ilvl')
        numId_el = numPr.find(f'{{{w_ns}}}numId')

        if numId_el is None:
            return (None, None)

        num_id = numId_el.get(f'{{{w_ns}}}val', '0')
        if num_id == '0':  # numId=0 means numbering removed
            return (None, None)

        ilvl = int(ilvl_el.get(f'{{{w_ns}}}val', '0')) if ilvl_el is not None else 0

        # Look up abstract definition
        abstract_id = num_to_abstract.get(num_id)
        if abstract_id is None or abstract_id not in abstract_nums:
            return (None, None)

        levels = abstract_nums[abstract_id]
        if ilvl not in levels:
            return (None, None)

        level_def = levels[ilvl]
        num_fmt = level_def['numFmt']
        lvl_text = level_def['lvlText']

        # Determine start value (override takes precedence)
        start_val = level_def['start']
        if num_id in num_overrides and ilvl in num_overrides[num_id]:
            start_val = num_overrides[num_id][ilvl]

        # ── Update counters ──
        if num_id not in counters:
            counters[num_id] = {}

        c = counters[num_id]
        prev_ilvl = last_ilvl_for_num.get(num_id)

        if prev_ilvl is None:
            # First item of this list
            c[ilvl] = start_val
        elif ilvl > prev_ilvl:
            # Going deeper → start new sub-numbering
            sub_start = levels.get(ilvl, level_def).get('start', 1)
            if num_id in num_overrides and ilvl in num_overrides[num_id]:
                sub_start = num_overrides[num_id][ilvl]
            c[ilvl] = sub_start
        elif ilvl == prev_ilvl:
            # Same level → increment
            c[ilvl] = c.get(ilvl, start_val - 1) + 1
        else:
            # Returning to shallower level → increment and reset deeper
            c[ilvl] = c.get(ilvl, start_val - 1) + 1
            for k in list(c.keys()):
                if k > ilvl:
                    del c[k]

        last_ilvl_for_num[num_id] = ilvl

        # ── Build prefix from lvlText template ──
        # lvlText examples: "%1.", "%1.%2", "(%1)", "%1)", "第%1章"
        if num_fmt == 'bullet':
            prefix = '•'
        else:
            prefix = lvl_text
            for li in range(10):  # max 10 nesting levels
                placeholder = f'%{li + 1}'
                if placeholder in prefix:
                    ldef = levels.get(li, level_def)
                    lfmt = ldef.get('numFmt', 'decimal')
                    value = c.get(li, ldef.get('start', 1))
                    prefix = prefix.replace(placeholder, _format_number(value, lfmt))

        # Ensure trailing space for readability
        if prefix and not prefix.endswith(' '):
            prefix += ' '

        return (prefix, ilvl)

    return resolve


def strip_numbering_prefix(text, numbering_prefix=None):
    """Strip auto-numbering prefix from translated text for docx rebuild.

    During extraction we prepend the auto-numbering prefix (e.g., "1. ", "a) ")
    to paragraph text for display and translation. Before writing back to the
    original docx (which retains its numbering XML), we must remove it to
    avoid double-numbering.

    Args:
        text: Translated text that may start with a numbering prefix.
        numbering_prefix: The exact prefix that was prepended (from metadata).

    Returns:
        str: Text with numbering prefix stripped.
    """
    if not text:
        return text

    stripped = text.lstrip()

    # Try exact prefix match first (most reliable)
    if numbering_prefix:
        clean_prefix = numbering_prefix.rstrip()
        if stripped.startswith(clean_prefix):
            after = stripped[len(clean_prefix):].lstrip()
            if after or not stripped[len(clean_prefix):]:
                return after

    # Fallback: pattern-based stripping for common numbering formats
    patterns = [
        r'^[(（][\d\w一二三四五六七八九十]+[)）][.、\s]*',     # (1) / （一）
        r'^[\d]+(?:\.[\d]+)*[.、)）]\s*',                       # 1. / 1.1. / 1) / 1、
        r'^第[\d一二三四五六七八九十百千]+[章节条款项]\s*',     # 第一章 / 第3节
        r'^[a-zA-Z][.)]\s*',                                     # a. / a) / A.
        r'^[IVXLCDM]+[.)]\s+',                                   # IV. / I)
        r'^[ivxlcdm]+[.)]\s+',                                   # iv. / i)
        r'^[•·●○■□▪▸►–—]\s*',                                   # bullets
        r'^[一二三四五六七八九十百千]+[、.）)]\s*',              # 一、/ 二、
    ]
    for pattern in patterns:
        new_text = re.sub(pattern, '', stripped, count=1)
        if new_text != stripped:
            return new_text

    return text


def extract_paragraphs_from_docx(
    file_content,
    image_mode="limited",
    max_image_bytes=2 * 1024 * 1024,
    max_total_image_bytes=16 * 1024 * 1024,
    max_images=80,
):
    """Extract paragraphs from a .docx file with style metadata.

    Args:
        file_content: Binary content of the .docx file.

    Returns:
        list[dict]: Each dict contains:
            - text (str): Paragraph text
            - style (str): Paragraph style name (e.g., 'Heading 1', 'Normal')
            - alignment (str|None): Alignment ('LEFT','CENTER','RIGHT','JUSTIFY')
            - bold (bool): Whether first run is bold
            - font_size (float|None): Font size in points
            - runs (list[dict]): Run-level detail for faithful reconstruction
    """
    if Document is None:
        raise ImportError("python-docx is required for .docx support. Install with: pip install python-docx")

    doc = Document(io.BytesIO(file_content))
    paragraphs = []

    # ── Namespace map for lxml queries ─────────────────────────────
    NSMAP = {
        'w':   'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r':   'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'v':   'urn:schemas-microsoft-com:vml',
        'mc':  'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup',
    }

    # ── Build image map from document-part relationships ───────────
    def _build_image_map(part):
        """Build rId -> base64 data URI map from a document part's relationships."""
        img_map = {}
        if not include_images:
            return img_map
        try:
            for rel_id, rel in part.rels.items():
                if "image" in rel.reltype:
                    try:
                        img_blob = rel.target_part.blob
                        img_size = len(img_blob or b"")
                        if (
                            not unlimited_images
                            and (
                                image_budget["count"] >= max_images
                                or img_size > max_image_bytes
                                or image_budget["bytes"] + img_size > max_total_image_bytes
                            )
                        ):
                            _logger.info(
                                "Skipping extracted DOCX image %s (%s bytes) to keep extraction lightweight",
                                rel_id,
                                img_size,
                            )
                            continue
                        ct = rel.target_part.content_type or "image/png"
                        img_b64 = base64.b64encode(img_blob).decode("ascii")
                        img_map[rel_id] = f"data:{ct};base64,{img_b64}"
                        image_budget["count"] += 1
                        image_budget["bytes"] += img_size
                    except Exception as e:
                        _logger.warning("Failed to extract image %s: %s", rel_id, e)
        except Exception as e:
            _logger.warning("Could not build image map: %s", e)
        return img_map

    include_images = image_mode != "none"
    unlimited_images = image_mode == "full"
    image_budget = {
        "count": 0,
        "bytes": 0,
    }

    # Body document image map
    body_image_map = _build_image_map(doc.part)

    # ── Helper: extract images from an XML element using a given image map ──
    def _extract_images_from_element(element, image_map):
        """Extract all <w:drawing> and <v:imagedata> images from an lxml element.

        Returns list of dicts with:
            data_uri, width, height,
            placement ('inline'|'anchor'),
            wrap_type ('none'|'square'|'tight'|'topAndBottom'|'through'|None),
            behind_doc (bool),
            align_h (str|None) – 'left'|'center'|'right',
            offset_h (int|None) – horizontal offset in px,
            offset_v (int|None) – vertical offset in px.
        """
        images = []
        try:
            # Modern drawings: <w:drawing>/<wp:inline> or <w:drawing>/<wp:anchor>
            for drawing in element.iter('{%s}drawing' % NSMAP['w']):
                blip = drawing.find('.//{%s}blip' % NSMAP['a'])
                if blip is None:
                    continue
                embed = blip.get('{%s}embed' % NSMAP['r'])
                if not embed or embed not in image_map:
                    continue

                # Determine if inline or anchor
                inline_el = drawing.find('{%s}inline' % NSMAP['wp'])
                anchor_el = drawing.find('{%s}anchor' % NSMAP['wp'])
                container = inline_el if inline_el is not None else anchor_el
                placement = 'inline' if inline_el is not None else 'anchor'

                # Size from <wp:extent>
                width_px, height_px = None, None
                if container is not None:
                    extent = container.find('{%s}extent' % NSMAP['wp'])
                    if extent is not None:
                        cx = extent.get('cx')
                        cy = extent.get('cy')
                        if cx:
                            width_px = round(int(cx) / 914400 * 96)
                        if cy:
                            height_px = round(int(cy) / 914400 * 96)

                # Anchor-specific: position, wrap, behind
                wrap_type = None
                behind_doc = False
                align_h = None
                offset_h = None
                offset_v = None

                if anchor_el is not None:
                    behind_doc = anchor_el.get('behindDoc', '0') == '1'

                    # Wrap type
                    wrap_tags = {
                        'wrapNone': 'none',
                        'wrapSquare': 'square',
                        'wrapTight': 'tight',
                        'wrapTopAndBottom': 'topAndBottom',
                        'wrapThrough': 'through',
                    }
                    for tag_local, wt in wrap_tags.items():
                        if anchor_el.find('{%s}%s' % (NSMAP['wp'], tag_local)) is not None:
                            wrap_type = wt
                            break

                    # Horizontal position
                    posH = anchor_el.find('{%s}positionH' % NSMAP['wp'])
                    if posH is not None:
                        align_el = posH.find('{%s}align' % NSMAP['wp'])
                        if align_el is not None and align_el.text:
                            align_h = align_el.text.strip().lower()  # left/center/right
                        offset_el = posH.find('{%s}posOffset' % NSMAP['wp'])
                        if offset_el is not None and offset_el.text:
                            offset_h = round(int(offset_el.text) / 914400 * 96)

                    # Vertical position
                    posV = anchor_el.find('{%s}positionV' % NSMAP['wp'])
                    if posV is not None:
                        offset_el = posV.find('{%s}posOffset' % NSMAP['wp'])
                        if offset_el is not None and offset_el.text:
                            offset_v = round(int(offset_el.text) / 914400 * 96)

                images.append({
                    "data_uri": image_map[embed],
                    "width": width_px,
                    "height": height_px,
                    "placement": placement,
                    "wrap_type": wrap_type,
                    "behind_doc": behind_doc,
                    "align_h": align_h,
                    "offset_h": offset_h,
                    "offset_v": offset_v,
                })

            # Legacy VML images: <v:imagedata r:id="rId..."/>
            for imgdata in element.iter('{%s}imagedata' % NSMAP['v']):
                embed = imgdata.get('{%s}id' % NSMAP['r'])
                if not embed or embed not in image_map:
                    continue
                # Try to get size from parent <v:shape style="width:...;height:...">
                width_px, height_px = None, None
                position_type = 'inline'
                vml_float = None
                parent_shape = imgdata.getparent()
                if parent_shape is not None:
                    style_attr = parent_shape.get('style', '')
                    w_match = re.search(r'width:\s*([\d.]+)pt', style_attr)
                    h_match = re.search(r'height:\s*([\d.]+)pt', style_attr)
                    if w_match:
                        width_px = round(float(w_match.group(1)) * 96 / 72)
                    if h_match:
                        height_px = round(float(h_match.group(1)) * 96 / 72)
                    if 'position:absolute' in style_attr:
                        position_type = 'anchor'
                    ml_match = re.search(r'mso-position-horizontal:\s*(\w+)', style_attr)
                    if ml_match:
                        vml_float = ml_match.group(1).lower()
                images.append({
                    "data_uri": image_map[embed],
                    "width": width_px,
                    "height": height_px,
                    "placement": position_type,
                    "wrap_type": None,
                    "behind_doc": False,
                    "align_h": vml_float,
                    "offset_h": None,
                    "offset_v": None,
                })
        except Exception as e:
            _logger.warning("Failed to extract images from element: %s", e)
        return images

    # ── Helper: check if an element is inside <mc:Fallback> ───────
    def _is_in_fallback(elem):
        """Check if an element is inside <mc:Fallback> (duplicate content)."""
        parent = elem.getparent()
        while parent is not None:
            if parent.tag == '{%s}Fallback' % NSMAP['mc']:
                return True
            parent = parent.getparent()
        return False

    # ── Helper: extract formatted paragraphs from <w:txbxContent> ─
    def _extract_textbox_paragraphs(txbx_content):
        """Extract paragraphs with formatting from a <w:txbxContent> element.

        Returns list of dicts: {text, bold, font_size, font_name, color, italic, alignment}
        """
        tb_paras = []
        for p in txbx_content.findall('{%s}p' % NSMAP['w']):
            # Collect all <w:t> text in this textbox paragraph
            # Only from direct <w:r> children, not nested textboxes
            runs = p.findall('{%s}r' % NSMAP['w'])
            para_text = ""
            rPr_bold = False
            rPr_font_size = None
            rPr_font_name = None
            rPr_color = None
            rPr_italic = False

            for r in runs:
                # Collect text from <w:t> direct children of this run
                for t in r.findall('{%s}t' % NSMAP['w']):
                    para_text += (t.text or "")
                # Get formatting from first run that has it
                if not rPr_font_size:
                    rPr = r.find('{%s}rPr' % NSMAP['w'])
                    if rPr is not None:
                        if rPr.find('{%s}b' % NSMAP['w']) is not None:
                            rPr_bold = True
                        sz = rPr.find('{%s}sz' % NSMAP['w'])
                        if sz is not None:
                            val = sz.get('{%s}val' % NSMAP['w'])
                            if val:
                                rPr_font_size = int(val) / 2  # half-points → pt
                        rFonts = rPr.find('{%s}rFonts' % NSMAP['w'])
                        if rFonts is not None:
                            rPr_font_name = (
                                rFonts.get('{%s}eastAsia' % NSMAP['w'])
                                or rFonts.get('{%s}ascii' % NSMAP['w'])
                            )
                        color_el = rPr.find('{%s}color' % NSMAP['w'])
                        if color_el is not None:
                            rPr_color = color_el.get('{%s}val' % NSMAP['w'])
                        if rPr.find('{%s}i' % NSMAP['w']) is not None:
                            rPr_italic = True

            # Get paragraph alignment
            alignment = None
            pPr = p.find('{%s}pPr' % NSMAP['w'])
            if pPr is not None:
                jc = pPr.find('{%s}jc' % NSMAP['w'])
                if jc is not None:
                    alignment = jc.get('{%s}val' % NSMAP['w'])

            tb_paras.append({
                'text': para_text,
                'bold': rPr_bold,
                'font_size': rPr_font_size,
                'font_name': rPr_font_name,
                'color': rPr_color,
                'italic': rPr_italic,
                'alignment': alignment,
            })
        return tb_paras

    # ── Helper: get textbox geometry from ancestors ────────────────
    def _get_textbox_geometry(txbx_content):
        """Walk up from <w:txbxContent> to find textbox dimensions and position.

        Returns dict with width, height (px), position_h, position_v (px),
        has_border, has_fill, fill_color, border_color.
        """
        result = {
            'width': None, 'height': None,
            'position_h': None, 'position_v': None,
            'has_border': True, 'has_fill': False,
            'fill_color': None, 'border_color': None,
        }
        parent = txbx_content.getparent()
        while parent is not None:
            tag = parent.tag

            # Modern: wp:anchor or wp:inline
            if tag in ('{%s}anchor' % NSMAP['wp'], '{%s}inline' % NSMAP['wp']):
                extent = parent.find('{%s}extent' % NSMAP['wp'])
                if extent is not None:
                    cx = extent.get('cx')
                    cy = extent.get('cy')
                    if cx:
                        result['width'] = round(int(cx) / 914400 * 96)
                    if cy:
                        result['height'] = round(int(cy) / 914400 * 96)
                # Position (only for anchored)
                posH = parent.find('{%s}positionH' % NSMAP['wp'])
                posV = parent.find('{%s}positionV' % NSMAP['wp'])
                if posH is not None:
                    offset = posH.find('{%s}posOffset' % NSMAP['wp'])
                    if offset is not None and offset.text:
                        result['position_h'] = round(int(offset.text) / 914400 * 96)
                if posV is not None:
                    offset = posV.find('{%s}posOffset' % NSMAP['wp'])
                    if offset is not None and offset.text:
                        result['position_v'] = round(int(offset.text) / 914400 * 96)
                # Border/fill from wps:spPr
                spPr = parent.find('.//{%s}spPr' % NSMAP['wps'])
                if spPr is not None:
                    no_fill = spPr.find('{%s}noFill' % NSMAP['a'])
                    solid_fill = spPr.find('.//{%s}solidFill' % NSMAP['a'])
                    if solid_fill is not None:
                        result['has_fill'] = True
                        srgb = solid_fill.find('{%s}srgbClr' % NSMAP['a'])
                        if srgb is not None:
                            result['fill_color'] = srgb.get('val')
                    elif no_fill is not None:
                        result['has_fill'] = False
                    ln = spPr.find('{%s}ln' % NSMAP['a'])
                    if ln is not None:
                        no_fill_ln = ln.find('{%s}noFill' % NSMAP['a'])
                        if no_fill_ln is not None:
                            result['has_border'] = False
                        else:
                            srgb = ln.find('.//{%s}srgbClr' % NSMAP['a'])
                            if srgb is not None:
                                result['border_color'] = srgb.get('val')
                break

            # VML: v:shape
            if tag == '{%s}shape' % NSMAP['v']:
                style_attr = parent.get('style', '')
                w_match = re.search(r'width:\s*([\d.]+)pt', style_attr)
                h_match = re.search(r'height:\s*([\d.]+)pt', style_attr)
                if w_match:
                    result['width'] = round(float(w_match.group(1)) * 96 / 72)
                if h_match:
                    result['height'] = round(float(h_match.group(1)) * 96 / 72)
                ml_match = re.search(r'margin-left:\s*([\d.-]+)pt', style_attr)
                mt_match = re.search(r'margin-top:\s*([\d.-]+)pt', style_attr)
                if ml_match:
                    result['position_h'] = round(float(ml_match.group(1)) * 96 / 72)
                if mt_match:
                    result['position_v'] = round(float(mt_match.group(1)) * 96 / 72)
                # Border
                stroked = parent.get('stroked', '')
                if stroked.lower() in ('f', 'false'):
                    result['has_border'] = False
                result['border_color'] = parent.get('strokecolor')
                # Fill
                filled = parent.get('filled', '')
                if filled.lower() in ('f', 'false'):
                    result['has_fill'] = False
                else:
                    fc = parent.get('fillcolor')
                    if fc:
                        result['has_fill'] = True
                        result['fill_color'] = fc
                break

            parent = parent.getparent()
        return result

    # ── Helper: extract text boxes from a paragraph element ────────
    def _extract_textboxes_from_para(para_element):
        """Extract structured textbox data from a paragraph element.

        Text boxes appear inside <mc:AlternateContent> or <w:pict>/<v:textbox>.
        Skips duplicates in <mc:Fallback> to avoid double-counting.

        Returns list of dicts with:
            - paragraphs: list of {text, bold, font_size, ...}
            - full_text: concatenated text with newlines
            - width, height, position_h, position_v (pixels)
            - has_border, has_fill, fill_color, border_color
        """
        textboxes = []
        seen_ids = set()
        try:
            for txbx_content in para_element.iter('{%s}txbxContent' % NSMAP['w']):
                # Skip duplicates (mc:Fallback contains a copy of mc:Choice)
                if _is_in_fallback(txbx_content):
                    continue
                if id(txbx_content) in seen_ids:
                    continue
                seen_ids.add(id(txbx_content))

                # Extract paragraphs with formatting
                tb_paras = _extract_textbox_paragraphs(txbx_content)
                full_text = "\n".join(p['text'] for p in tb_paras)

                if not full_text.strip():
                    continue

                # Get geometry (dimensions, position, border/fill)
                geo = _get_textbox_geometry(txbx_content)

                textboxes.append({
                    'paragraphs': tb_paras,
                    'full_text': full_text,
                    'width': geo['width'],
                    'height': geo['height'],
                    'position_h': geo['position_h'],
                    'position_v': geo['position_v'],
                    'has_border': geo['has_border'],
                    'has_fill': geo['has_fill'],
                    'fill_color': geo['fill_color'],
                    'border_color': geo['border_color'],
                })
        except Exception as e:
            _logger.warning("Failed to extract textboxes: %s", e)
        return textboxes

    # ── Extract header / footer (text + images) ───────────────────
    def _paragraph_current_text(para):
        text_parts = []

        def walk(parent):
            for child in parent:
                if child.tag == '{%s}t' % NSMAP['w']:
                    text_parts.append(child.text or "")
                elif child.tag == '{%s}tab' % NSMAP['w']:
                    text_parts.append("\t")
                elif child.tag == '{%s}br' % NSMAP['w']:
                    text_parts.append("\n")
                elif child.tag == '{%s}del' % NSMAP['w']:
                    continue
                else:
                    walk(child)

        walk(para._element)
        text = "".join(text_parts)
        return text if text else para.text

    header_text = ""
    header_images = []
    footer_text = ""
    footer_images = []
    try:
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                hdr_parts = [
                    _paragraph_current_text(p)
                    for p in section.header.paragraphs
                    if _paragraph_current_text(p).strip()
                ]
                if hdr_parts or not header_text:
                    header_text = "\n".join(hdr_parts) if hdr_parts else ""
                # Extract images from header part relationships
                if not header_images:
                    try:
                        hdr_image_map = _build_image_map(section.header.part)
                        header_images = _extract_images_from_element(
                            section.header._element, hdr_image_map
                        )
                    except Exception:
                        # header.part may not exist; try body image map
                        header_images = _extract_images_from_element(
                            section.header._element, body_image_map
                        )
                if header_text or header_images:
                    break
        for section in doc.sections:
            if section.footer and section.footer.paragraphs:
                ftr_parts = [
                    _paragraph_current_text(p)
                    for p in section.footer.paragraphs
                    if _paragraph_current_text(p).strip()
                ]
                if ftr_parts or not footer_text:
                    footer_text = "\n".join(ftr_parts) if ftr_parts else ""
                if not footer_images:
                    try:
                        ftr_image_map = _build_image_map(section.footer.part)
                        footer_images = _extract_images_from_element(
                            section.footer._element, ftr_image_map
                        )
                    except Exception:
                        footer_images = _extract_images_from_element(
                            section.footer._element, body_image_map
                        )
                if footer_text or footer_images:
                    break
    except Exception as e:
        _logger.warning("Could not extract header/footer: %s", e)

    # ── Helper: get font size from paragraph style hierarchy ──────
    def _get_style_font_size(para):
        try:
            style = para.style
            while style:
                if style.font and style.font.size:
                    return style.font.size.pt
                style = style.base_style
        except Exception:
            pass
        return None

    # ── Build numbering engine for auto-numbered lists ─────────────
    def _extract_run_info_from_xml(run_el, style_font_size, revision_type=None):
        text_parts = []
        for child in run_el:
            if child.tag == '{%s}t' % NSMAP['w']:
                text_parts.append(child.text or "")
            elif child.tag == '{%s}tab' % NSMAP['w']:
                text_parts.append("\t")
            elif child.tag == '{%s}br' % NSMAP['w']:
                text_parts.append("\n")
        text = "".join(text_parts)
        if not text:
            return None

        run_info = {
            "text": text,
            "bold": False,
            "italic": False,
            "underline": False,
        }
        if revision_type:
            run_info["revision_type"] = revision_type
            if revision_type == "inserted":
                run_info["inserted"] = True

        rPr = run_el.find('{%s}rPr' % NSMAP['w'])
        if rPr is not None:
            run_info["bold"] = rPr.find('{%s}b' % NSMAP['w']) is not None
            run_info["italic"] = rPr.find('{%s}i' % NSMAP['w']) is not None
            u_el = rPr.find('{%s}u' % NSMAP['w'])
            run_info["underline"] = bool(
                u_el is not None and
                u_el.get('{%s}val' % NSMAP['w'], 'single') != 'none'
            )
            sz = rPr.find('{%s}sz' % NSMAP['w'])
            if sz is not None:
                val = sz.get('{%s}val' % NSMAP['w'])
                if val:
                    try:
                        run_info["font_size"] = int(val) / 2
                    except (ValueError, TypeError):
                        pass
            color_el = rPr.find('{%s}color' % NSMAP['w'])
            if color_el is not None:
                color = color_el.get('{%s}val' % NSMAP['w'])
                if color and color.lower() != "auto":
                    run_info["color"] = color
            rFonts = rPr.find('{%s}rFonts' % NSMAP['w'])
            if rFonts is not None:
                font_name = (
                    rFonts.get('{%s}eastAsia' % NSMAP['w'])
                    or rFonts.get('{%s}ascii' % NSMAP['w'])
                )
                if font_name:
                    run_info["font_name"] = font_name

        if "font_size" not in run_info and style_font_size:
            run_info["font_size"] = style_font_size
        return run_info

    def _extract_revision_aware_runs(para, style_font_size):
        runs_data = []

        def walk(parent, revision_type=None):
            for child in parent:
                if child.tag == '{%s}r' % NSMAP['w']:
                    run_info = _extract_run_info_from_xml(
                        child, style_font_size, revision_type
                    )
                    if run_info:
                        runs_data.append(run_info)
                elif child.tag == '{%s}ins' % NSMAP['w']:
                    walk(child, "inserted")
                elif child.tag == '{%s}del' % NSMAP['w']:
                    continue
                elif child.tag in (
                    '{%s}hyperlink' % NSMAP['w'],
                    '{%s}smartTag' % NSMAP['w'],
                    '{%s}sdt' % NSMAP['w'],
                ):
                    walk(child, revision_type)

        walk(para._element)
        text = "".join(run.get("text", "") for run in runs_data)
        first_run = runs_data[0] if runs_data else {}
        return (
            text,
            runs_data,
            bool(first_run.get("bold")),
            first_run.get("font_size") or style_font_size,
        )

    _resolve_numbering = _build_numbering_engine(doc)

    # ── Helper: extract runs data from a paragraph ──────────────
    def _extract_runs_data(para, style_font_size):
        """Extract run-level formatting from a python-docx Paragraph."""
        _text, revision_runs, first_bold, first_font_size = (
            _extract_revision_aware_runs(para, style_font_size)
        )
        if revision_runs:
            return revision_runs, first_bold, first_font_size

        runs_data = []
        first_bold = False
        first_font_size = None
        for i, run in enumerate(para.runs):
            run_info = {
                "text": run.text,
                "bold": run.bold or False,
                "italic": run.italic or False,
                "underline": run.underline or False,
            }
            if run.font.size:
                run_info["font_size"] = run.font.size.pt
            elif style_font_size:
                run_info["font_size"] = style_font_size
            try:
                if run.font.color and run.font.color.rgb:
                    run_info["color"] = str(run.font.color.rgb)
            except (ValueError, KeyError, AttributeError):
                pass
            if run.font.name:
                run_info["font_name"] = run.font.name
            runs_data.append(run_info)

            if i == 0:
                first_bold = run.bold or False
                if run.font.size:
                    first_font_size = run.font.size.pt
                elif style_font_size:
                    first_font_size = style_font_size
        if first_font_size is None and style_font_size:
            first_font_size = style_font_size
        return runs_data, first_bold, first_font_size

    # ── Helper: check if a table cell is a vertical-merge continuation ──
    def _is_vmerge_continuation(tc_element):
        """Return True if <w:tc> is a vertical merge continuation (not restart)."""
        tcPr = tc_element.find('{%s}tcPr' % NSMAP['w'])
        if tcPr is None:
            return False
        vMerge = tcPr.find('{%s}vMerge' % NSMAP['w'])
        if vMerge is None:
            return False
        # <w:vMerge val="restart"/> starts a new merged group — NOT a continuation
        # <w:vMerge/> (no val) or val="" is a continuation — skip it
        val = vMerge.get('{%s}val' % NSMAP['w'])
        return val != 'restart'

    # ── Process body elements (paragraphs AND tables) in document order ──
    # doc.paragraphs only returns body-level <w:p> elements, completely
    # missing <w:tbl> (table) elements.  We iterate doc.element.body
    # children to get both in proper document order.
    all_body_paras = list(doc.paragraphs)
    all_tables = list(doc.tables)
    para_counter = 0
    table_counter = 0
    w_p_tag = '{%s}p' % NSMAP['w']
    w_tbl_tag = '{%s}tbl' % NSMAP['w']

    for child in doc.element.body:
        # ── Body-level paragraph ──────────────────────────────────
        if child.tag == w_p_tag:
            if para_counter >= len(all_body_paras):
                para_counter += 1
                continue
            para = all_body_paras[para_counter]
            para_idx = para_counter
            para_counter += 1

            style_font_size = _get_style_font_size(para)
            text, runs_data, first_bold, first_font_size = (
                _extract_revision_aware_runs(para, style_font_size)
            )
            if not runs_data:
                text = para.text
                runs_data, first_bold, first_font_size = _extract_runs_data(
                    para, style_font_size
                )

            # Resolve auto-numbering prefix
            numbering_prefix, numbering_level = _resolve_numbering(
                para._element, para.style
            )
            if numbering_prefix:
                text = numbering_prefix + text

            # Extract images
            para_images = _extract_images_from_element(
                para._element, body_image_map
            )

            # Extract text box content
            textbox_data = _extract_textboxes_from_para(para._element)

            if not text.strip() and not para_images and not textbox_data:
                paragraphs.append({
                    "text": "",
                    "style": para.style.name if para.style else "Normal",
                    "alignment": None,
                    "bold": False,
                    "font_size": None,
                    "runs": [],
                    "images": [],
                    "is_empty": True,
                    "para_index": para_idx,
                    "numbering_prefix": numbering_prefix,
                    "numbering_level": numbering_level,
                })
                continue

            alignment = None
            try:
                if para.alignment is not None:
                    alignment_map = {
                        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
                        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
                        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
                        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
                    }
                    alignment = alignment_map.get(para.alignment, None)
            except (ValueError, KeyError):
                # Some docs use non-standard alignment values (e.g. 'start', 'end')
                # that python-docx cannot map. Fall back to reading raw XML.
                _w = NSMAP['w']
                pPr = para._element.find(f'{{{_w}}}pPr')
                if pPr is not None:
                    jc = pPr.find(f'{{{_w}}}jc')
                    if jc is not None:
                        raw = jc.get(f'{{{_w}}}val', '')
                        raw_map = {
                            'left': 'LEFT', 'start': 'LEFT',
                            'center': 'CENTER',
                            'right': 'RIGHT', 'end': 'RIGHT',
                            'both': 'JUSTIFY', 'justify': 'JUSTIFY',
                        }
                        alignment = raw_map.get(raw.lower())

            # Build textboxes array
            textboxes_meta = []
            for tb_idx, tb in enumerate(textbox_data):
                textboxes_meta.append({
                    "full_text": tb['full_text'],
                    "textbox_index": tb_idx,
                    "paragraphs": tb['paragraphs'],
                    "width": tb['width'],
                    "height": tb['height'],
                    "position_h": tb['position_h'],
                    "position_v": tb['position_v'],
                    "has_border": tb['has_border'],
                    "has_fill": tb['has_fill'],
                    "fill_color": tb['fill_color'],
                    "border_color": tb['border_color'],
                })

            paragraphs.append({
                "text": text,
                "style": para.style.name if para.style else "Normal",
                "alignment": alignment,
                "bold": first_bold,
                "font_size": first_font_size,
                "runs": runs_data,
                "images": para_images,
                "textboxes": textboxes_meta,
                "is_empty": (
                    not text.strip() and not para_images and not textboxes_meta
                ),
                "para_index": para_idx,
                "numbering_prefix": numbering_prefix,
                "numbering_level": numbering_level,
            })

        # ── Body-level table ──────────────────────────────────────
        elif child.tag == w_tbl_tag:
            if table_counter >= len(all_tables):
                table_counter += 1
                continue
            table = all_tables[table_counter]
            tbl_idx = table_counter
            table_counter += 1

            total_rows = len(table.rows)
            try:
                total_cols = len(table.columns)
            except Exception:
                total_cols = 0

            # ── Pre-scan: compute rowspan for vertically-merged cells ──
            # vmerge_spans[(row_idx, grid_col)] = rowspan count
            # We scan each column: when we see vMerge restart, count
            # how many continuation rows follow.
            w_ns = NSMAP['w']
            vmerge_spans = {}
            for pre_row_idx, pre_row in enumerate(table.rows):
                pre_seen = set()
                pre_grid_col = 0
                for pre_cell in pre_row.cells:
                    pre_tc = pre_cell._tc
                    if id(pre_tc) in pre_seen:
                        continue
                    pre_seen.add(id(pre_tc))

                    pre_gs = 1
                    pre_tcPr = pre_tc.find('{%s}tcPr' % w_ns)
                    if pre_tcPr is not None:
                        gs_el = pre_tcPr.find('{%s}gridSpan' % w_ns)
                        if gs_el is not None:
                            try:
                                pre_gs = int(gs_el.get('{%s}val' % w_ns, '1'))
                            except (ValueError, TypeError):
                                pre_gs = 1

                    cur_gc = pre_grid_col
                    pre_grid_col += pre_gs

                    # Check if this cell starts a vMerge group
                    if pre_tcPr is not None:
                        vm = pre_tcPr.find('{%s}vMerge' % w_ns)
                        if vm is not None and vm.get('{%s}val' % w_ns) == 'restart':
                            # Count continuation rows below
                            span = 1
                            for next_ri in range(pre_row_idx + 1, total_rows):
                                next_row = table.rows[next_ri]
                                ns2 = set()
                                ngc = 0
                                found_cont = False
                                for nc in next_row.cells:
                                    ntc = nc._tc
                                    if id(ntc) in ns2:
                                        continue
                                    ns2.add(id(ntc))
                                    ngs = 1
                                    ntcPr = ntc.find('{%s}tcPr' % w_ns)
                                    if ntcPr is not None:
                                        ngs_el = ntcPr.find('{%s}gridSpan' % w_ns)
                                        if ngs_el is not None:
                                            try:
                                                ngs = int(ngs_el.get('{%s}val' % w_ns, '1'))
                                            except (ValueError, TypeError):
                                                ngs = 1
                                    if ngc == cur_gc:
                                        # Same grid column — check if continuation
                                        if ntcPr is not None:
                                            nvm = ntcPr.find('{%s}vMerge' % w_ns)
                                            if nvm is not None and nvm.get('{%s}val' % w_ns) != 'restart':
                                                span += 1
                                                found_cont = True
                                        break
                                    ngc += ngs
                                if not found_cont:
                                    break
                            if span > 1:
                                vmerge_spans[(pre_row_idx, cur_gc)] = span

            # Collect cells per row, then emit ONE paragraph per row
            # with cells joined by [CELL] separator.
            for row_idx, row in enumerate(table.rows):
                row_seen = set()   # per-row dedup (horizontal merge)
                row_cells = []     # list of {text, col_index, runs, bold, font_size}

                grid_col = 0  # actual grid column position
                for cell in row.cells:
                    tc = cell._tc
                    if id(tc) in row_seen:
                        continue
                    row_seen.add(id(tc))

                    # Read gridSpan (horizontal merge width)
                    grid_span = 1
                    tcPr = tc.find('{%s}tcPr' % NSMAP['w'])
                    if tcPr is not None:
                        gs = tcPr.find('{%s}gridSpan' % NSMAP['w'])
                        if gs is not None:
                            try:
                                grid_span = int(gs.get('{%s}val' % NSMAP['w'], '1'))
                            except (ValueError, TypeError):
                                grid_span = 1

                    cur_col = grid_col
                    grid_col += grid_span

                    # Skip vertical-merge continuation cells
                    if _is_vmerge_continuation(tc):
                        continue

                    # Collect text from all paragraphs in this cell
                    cell_text_parts = []
                    cell_runs = []
                    cell_bold = False
                    cell_font_size = None
                    first_run_seen = False

                    for cp in cell.paragraphs:
                        cp_font_size = _get_style_font_size(cp)
                        cp_text, cp_runs, cp_bold, cp_first_font_size = (
                            _extract_revision_aware_runs(cp, cp_font_size)
                        )
                        if not cp_runs:
                            cp_text = cp.text
                            cp_runs, cp_bold, cp_first_font_size = _extract_runs_data(
                                cp, cp_font_size
                            )
                        cell_text_parts.append(cp_text)
                        cell_runs.extend(cp_runs)
                        if cp_runs and not first_run_seen:
                            cell_bold = cp_bold
                            cell_font_size = cp_first_font_size
                            first_run_seen = True

                    cell_text = " ".join(
                        line.strip() for line in cell_text_parts
                        if line.strip()
                    )
                    row_cells.append({
                        "text": cell_text,
                        "col_index": cur_col,
                        "grid_span": grid_span,
                        "row_span": vmerge_spans.get((row_idx, cur_col), 1),
                        "runs": cell_runs,
                        "bold": cell_bold,
                        "font_size": cell_font_size,
                    })

                # Skip fully empty rows
                if not any(c["text"].strip() for c in row_cells):
                    continue

                # Build combined row text: cell1 [CELL] cell2 [CELL] cell3
                row_text = " [CELL] ".join(c["text"] for c in row_cells)

                # Merge all runs across cells (for formatting info)
                all_runs = []
                first_bold = False
                first_font_size = None
                for ci, c in enumerate(row_cells):
                    all_runs.extend(c["runs"])
                    if ci == 0:
                        first_bold = c["bold"]
                        first_font_size = c["font_size"]

                # Build cells metadata for rebuild + frontend rendering
                cells_meta = []
                for c in row_cells:
                    cells_meta.append({
                        "col_index": c["col_index"],
                        "grid_span": c["grid_span"],
                        "row_span": c["row_span"],
                        "text": c["text"],
                        "bold": c["bold"],
                        "font_size": c["font_size"],
                        "runs": c["runs"],
                    })

                paragraphs.append({
                    "text": row_text,
                    "style": "Table Row",
                    "alignment": None,
                    "bold": first_bold,
                    "font_size": first_font_size,
                    "runs": all_runs,
                    "images": [],
                    "textboxes": [],
                    "is_empty": False,
                    "para_index": None,
                    "is_table_row": True,
                    "table_index": tbl_idx,
                    "row_index": row_idx,
                    "cells": cells_meta,
                    "table_row_count": total_rows,
                    "table_col_count": total_cols,
                    "numbering_prefix": None,
                    "numbering_level": None,
                })

    return {
        "paragraphs": paragraphs,
        "header_text": header_text,
        "footer_text": footer_text,
        "header_images": header_images,
        "footer_images": footer_images,
    }


def extract_paragraphs_from_doc(
    file_content,
    image_mode="limited",
    max_image_bytes=2 * 1024 * 1024,
    max_total_image_bytes=16 * 1024 * 1024,
    max_images=80,
):
    """Extract paragraphs from a .doc file by converting to .docx via LibreOffice.

    Uses LibreOffice headless mode to convert .doc → .docx, then delegates
    to extract_paragraphs_from_docx() for full style-preserving extraction.

    Args:
        file_content: Binary content of the .doc file.

    Returns:
        dict: Same format as extract_paragraphs_from_docx (paragraphs + header/footer).

    Raises:
        RuntimeError: If LibreOffice conversion fails.
    """
    docx_content = _convert_doc_to_docx_via_libreoffice(file_content)
    return extract_paragraphs_from_docx(
        docx_content,
        image_mode=image_mode,
        max_image_bytes=max_image_bytes,
        max_total_image_bytes=max_total_image_bytes,
        max_images=max_images,
    )


def _convert_doc_to_docx_via_libreoffice(file_content):
    """Convert a .doc file to .docx using LibreOffice headless.

    Writes the .doc to a temp directory, invokes LibreOffice to convert,
    then reads back the resulting .docx.

    Args:
        file_content: Binary content of the .doc file.

    Returns:
        bytes: Binary content of the converted .docx file.

    Raises:
        RuntimeError: If LibreOffice is not installed or conversion fails.
    """
    tmp_dir = tempfile.mkdtemp(prefix='llm_translate_')
    doc_path = os.path.join(tmp_dir, 'input.doc')

    try:
        # Write .doc to temp file
        with open(doc_path, 'wb') as f:
            f.write(file_content)

        # Run LibreOffice headless conversion
        _logger.info("Converting .doc to .docx via LibreOffice: %s", doc_path)
        result = subprocess.run(
            [
                'libreoffice',
                '--headless',
                '--norestore',
                '--convert-to', 'docx',
                '--outdir', tmp_dir,
                doc_path,
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )

        if result.returncode != 0:
            _logger.error("LibreOffice conversion failed (rc=%s): %s",
                          result.returncode, result.stderr)
            raise RuntimeError(
                f"LibreOffice conversion failed: {result.stderr or 'unknown error'}. "
                f"Please ensure LibreOffice is installed (apt-get install libreoffice)."
            )

        # Find the output .docx file
        docx_path = os.path.join(tmp_dir, 'input.docx')
        if not os.path.exists(docx_path):
            # LibreOffice sometimes uses different naming; scan for .docx
            docx_files = [f for f in os.listdir(tmp_dir) if f.endswith('.docx')]
            if docx_files:
                docx_path = os.path.join(tmp_dir, docx_files[0])
            else:
                raise RuntimeError(
                    "LibreOffice conversion produced no .docx output. "
                    f"stdout: {result.stdout}, stderr: {result.stderr}"
                )

        # Read converted file
        with open(docx_path, 'rb') as f:
            docx_content = f.read()

        _logger.info("LibreOffice conversion successful: %d bytes", len(docx_content))
        return docx_content

    except FileNotFoundError:
        raise RuntimeError(
            "LibreOffice not found. Install with: apt-get install libreoffice"
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError(
            "LibreOffice conversion timed out (120s). The file may be too large or corrupted."
        )
    finally:
        # Clean up temp files
        import shutil
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _convert_pdf_to_docx_via_libreoffice(file_content):
    """Convert a PDF file to .docx using LibreOffice headless.

    Args:
        file_content: Binary content of the PDF file.

    Returns:
        bytes: Binary content of the converted .docx file.

    Raises:
        RuntimeError: If LibreOffice is not installed or conversion fails.
    """
    tmp_dir = tempfile.mkdtemp(prefix='llm_translate_pdf_')
    pdf_path = os.path.join(tmp_dir, 'input.pdf')

    try:
        with open(pdf_path, 'wb') as f:
            f.write(file_content)

        _logger.info("Converting PDF to DOCX via LibreOffice: %s", pdf_path)
        result = subprocess.run(
            [
                'libreoffice',
                '--headless',
                '--norestore',
                '--infilter=writer_pdf_import',
                '--convert-to', 'docx',
                '--outdir', tmp_dir,
                pdf_path,
            ],
            capture_output=True,
            text=True,
            timeout=180,
        )

        if result.returncode != 0:
            _logger.error("LibreOffice PDF conversion failed (rc=%s): %s",
                          result.returncode, result.stderr)
            raise RuntimeError(
                f"LibreOffice PDF→DOCX conversion failed: {result.stderr or 'unknown error'}"
            )

        docx_path = os.path.join(tmp_dir, 'input.docx')
        if not os.path.exists(docx_path):
            docx_files = [f for f in os.listdir(tmp_dir) if f.endswith('.docx')]
            if docx_files:
                docx_path = os.path.join(tmp_dir, docx_files[0])
            else:
                raise RuntimeError(
                    "LibreOffice PDF conversion produced no .docx output. "
                    f"stdout: {result.stdout}, stderr: {result.stderr}"
                )

        with open(docx_path, 'rb') as f:
            docx_content = f.read()

        _logger.info("PDF→DOCX conversion successful: %d bytes", len(docx_content))
        return docx_content

    except FileNotFoundError:
        raise RuntimeError(
            "LibreOffice not found. Install with: apt-get install libreoffice"
        )
    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice PDF conversion timed out (180s).")
    finally:
        import shutil
        shutil.rmtree(tmp_dir, ignore_errors=True)


def rebuild_docx(rebuild_data):
    """Rebuild a .docx file from translated paragraph data (creates new doc).

    This is the fallback approach that creates a new document.
    Prefer rebuild_docx_from_original() for better fidelity.

    Args:
        rebuild_data: dict with keys:
            - paragraphs (list[dict]): Body paragraph data
            - header_text (str): Translated header text
            - footer_text (str): Translated footer text
          OR a plain list of dicts (backward compatible).

    Returns:
        bytes: Binary content of the new .docx file.
    """
    if Document is None:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    # Handle both old list format and new dict format
    if isinstance(rebuild_data, list):
        paragraphs_data = rebuild_data
        header_text = ""
        footer_text = ""
    else:
        paragraphs_data = rebuild_data.get("paragraphs", [])
        header_text = rebuild_data.get("header_text", "")
        footer_text = rebuild_data.get("footer_text", "")

    doc = Document()

    for para_data in paragraphs_data:
        text = para_data.get("translated_text", "")
        meta = para_data.get("style_metadata", {})

        # Extract only paragraph text (strip textbox portions from combined text)
        import re
        if "[TEXTBOX]" in text:
            text = re.split(r'\s*\[TEXTBOX\]\s*', text)[0].strip()

        if meta.get("is_empty"):
            doc.add_paragraph("")
            continue

        style_name = meta.get("style", "Normal")

        # Try to use the original style, fall back to Normal
        try:
            para = doc.add_paragraph(style=style_name)
        except Exception:
            para = doc.add_paragraph()

        # Set alignment
        alignment = meta.get("alignment")
        if alignment:
            alignment_map = {
                "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            if alignment in alignment_map:
                para.alignment = alignment_map[alignment]

        # Reconstruct runs if available, else plain text
        original_runs = meta.get("runs", [])
        if original_runs and text:
            # Simple approach: put all translated text in one run with first run's style
            run = para.add_run(text)
            first_run = original_runs[0]
            run.bold = first_run.get("bold", False)
            run.italic = first_run.get("italic", False)
            run.underline = first_run.get("underline", False)
            if first_run.get("font_size"):
                run.font.size = Pt(first_run["font_size"])
            if first_run.get("color"):
                try:
                    run.font.color.rgb = RGBColor.from_string(first_run["color"])
                except Exception:
                    pass
        else:
            run = para.add_run(text)
            if meta.get("bold"):
                run.bold = True
            if meta.get("font_size"):
                run.font.size = Pt(meta["font_size"])

    # Write translated header/footer to document sections
    if header_text or footer_text:
        for section in doc.sections:
            if header_text:
                hdr = section.header
                hdr.is_linked_to_previous = False
                for p in hdr.paragraphs:
                    p.clear()
                target_p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
                run = target_p.add_run(header_text)
                run.font.size = Pt(9)
                target_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if footer_text:
                ftr = section.footer
                ftr.is_linked_to_previous = False
                for p in ftr.paragraphs:
                    p.clear()
                target_p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
                run = target_p.add_run(footer_text)
                run.font.size = Pt(9)
                target_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def strip_formatting_tags(text):
    """Remove <b>/<i>/<u> formatting markers from translated text.

    The LLM preserves these tags for UI rendering, but they must be
    stripped before writing into the .docx (which has its own run formatting).

    Args:
        text: Translated text potentially containing <b>/<i>/<u> tags.

    Returns:
        str: Clean text without formatting tags.
    """
    if not text:
        return text
    return re.sub(r'</?[biu]>', '', text)


def _parse_formatting_segments(text):
    """Parse text with <b>/<i>/<u> tags into formatting segments.

    Returns a list of dicts: [{"text": "...", "bold": bool, "italic": bool, "underline": bool}, ...]
    Used by _replace_para_text to distribute text across runs with proper formatting.
    """
    if not text:
        return [{"text": "", "bold": False, "italic": False, "underline": False}]

    TAG_RE = re.compile(r'<(/?)([biu])>')

    segments = []
    active = {"b": False, "i": False, "u": False}
    last_idx = 0

    for m in TAG_RE.finditer(text):
        # Text before this tag
        if m.start() > last_idx:
            chunk = text[last_idx:m.start()]
            if chunk:
                segments.append({
                    "text": chunk,
                    "bold": active["b"],
                    "italic": active["i"],
                    "underline": active["u"],
                })
        is_close = m.group(1) == "/"
        tag = m.group(2)
        active[tag] = not is_close
        last_idx = m.end()

    # Remaining text after last tag
    if last_idx < len(text):
        chunk = text[last_idx:]
        if chunk:
            segments.append({
                "text": chunk,
                "bold": active["b"],
                "italic": active["i"],
                "underline": active["u"],
            })

    if not segments:
        # No tags found, return text as-is
        plain = re.sub(r'</?[biu]>', '', text)
        return [{"text": plain, "bold": False, "italic": False, "underline": False}]

    return segments


# =========================================================================
# Image OCR Text Box Helpers
# =========================================================================

# Conversion: 1 pixel (at 96 DPI) = 9525 EMU
_PX_TO_EMU = 9525


def _xml_escape(text):
    """Escape text for safe inclusion in XML content."""
    if not text:
        return ""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _make_textbox_xml(text, offset_h_emu, offset_v_emu, width_emu, height_emu,
                       font_size_hpt, doc_pr_id):
    """Create a <w:r> element containing a floating text box overlay.

    The text box is:
    - Positioned as a floating anchor (relativeFrom column/paragraph)
    - White opaque fill (covers original text in the image)
    - No border
    - Text centered vertically in the box
    - Wrapped in <mc:AlternateContent> for broad Word compatibility

    Returns:
        lxml Element: <w:r> element ready to append to a paragraph.
    """
    from lxml import etree

    # Ensure minimum dimensions
    width_emu = max(int(width_emu), 100000)
    height_emu = max(int(height_emu), 100000)
    offset_h_emu = max(0, int(offset_h_emu))
    offset_v_emu = max(0, int(offset_v_emu))

    escaped_text = _xml_escape(text)

    xml_str = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        '     xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'
        '     xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        '     xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"'
        '     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        '     xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
        '     xmlns:v="urn:schemas-microsoft-com:vml"'
        '     xmlns:o="urn:schemas-microsoft-com:office:office">'
        '  <mc:AlternateContent>'
        '    <mc:Choice Requires="wps">'
        '      <w:drawing>'
        '        <wp:anchor distT="0" distB="0" distL="0" distR="0"'
        '                   simplePos="0" relativeHeight="251660000"'
        '                   behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1"'
        f'                  wp14:anchorId="{doc_pr_id:08X}">'
        '          <wp:simplePos x="0" y="0"/>'
        '          <wp:positionH relativeFrom="column">'
        f'            <wp:posOffset>{offset_h_emu}</wp:posOffset>'
        '          </wp:positionH>'
        '          <wp:positionV relativeFrom="paragraph">'
        f'            <wp:posOffset>{offset_v_emu}</wp:posOffset>'
        '          </wp:positionV>'
        f'          <wp:extent cx="{width_emu}" cy="{height_emu}"/>'
        '          <wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '          <wp:wrapNone/>'
        f'          <wp:docPr id="{doc_pr_id}" name="OCR TextBox {doc_pr_id}"/>'
        '          <a:graphic>'
        '            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '              <wps:wsp>'
        '                <wps:cNvSpPr txBox="1"/>'
        '                <wps:spPr>'
        '                  <a:xfrm>'
        '                    <a:off x="0" y="0"/>'
        f'                    <a:ext cx="{width_emu}" cy="{height_emu}"/>'
        '                  </a:xfrm>'
        '                  <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '                  <a:solidFill><a:srgbClr val="FFFFFF"/></a:solidFill>'
        '                  <a:ln w="0"><a:noFill/></a:ln>'
        '                </wps:spPr>'
        '                <wps:txbx>'
        '                  <w:txbxContent>'
        '                    <w:p>'
        '                      <w:pPr>'
        '                        <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
        '                      </w:pPr>'
        '                      <w:r>'
        '                        <w:rPr>'
        f'                          <w:sz w:val="{font_size_hpt}"/>'
        f'                          <w:szCs w:val="{font_size_hpt}"/>'
        '                        </w:rPr>'
        f'                        <w:t xml:space="preserve">{escaped_text}</w:t>'
        '                      </w:r>'
        '                    </w:p>'
        '                  </w:txbxContent>'
        '                </wps:txbx>'
        '                <wps:bodyPr rot="0" vert="horz" wrap="square"'
        '                            lIns="36000" tIns="18000" rIns="36000" bIns="18000"'
        '                            anchor="ctr" anchorCtr="0">'
        '                  <a:noAutofit/>'
        '                </wps:bodyPr>'
        '              </wps:wsp>'
        '            </a:graphicData>'
        '          </a:graphic>'
        '        </wp:anchor>'
        '      </w:drawing>'
        '    </mc:Choice>'
        '    <mc:Fallback>'
        '      <w:pict>'
        f'        <v:shape id="OCRBox{doc_pr_id}" o:spid="_x0000_s{doc_pr_id}"'
        '                  type="#_x0000_t202"'
        f'                  style="position:absolute;margin-left:{round(offset_h_emu / 12700, 1)}pt;'
        f'margin-top:{round(offset_v_emu / 12700, 1)}pt;'
        f'width:{round(width_emu / 12700, 1)}pt;'
        f'height:{round(height_emu / 12700, 1)}pt;z-index:251660000"'
        '                  fillcolor="white" stroked="f">'
        '          <v:textbox inset="1mm,0.5mm,1mm,0.5mm">'
        '            <w:txbxContent>'
        '              <w:p>'
        '                <w:r>'
        '                  <w:rPr>'
        f'                    <w:sz w:val="{font_size_hpt}"/>'
        f'                    <w:szCs w:val="{font_size_hpt}"/>'
        '                  </w:rPr>'
        f'                  <w:t xml:space="preserve">{escaped_text}</w:t>'
        '                </w:r>'
        '              </w:p>'
        '            </w:txbxContent>'
        '          </v:textbox>'
        '        </v:shape>'
        '      </w:pict>'
        '    </mc:Fallback>'
        '  </mc:AlternateContent>'
        '</w:r>'
    )
    return etree.fromstring(xml_str)


def _insert_image_ocr_textboxes(doc, image_ocr_results):
    """Insert translated text boxes overlaying images in the document.

    For each OCR result, finds the actual image in the docx XML to read
    its display size in EMU, then creates floating text boxes positioned
    precisely over the detected text blocks.

    Args:
        doc: python-docx Document object (already opened, will be modified).
        image_ocr_results: list of dicts from _prepare_rebuild_data.
    """
    if not image_ocr_results:
        return

    NSMAP = {
        'w':  'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a':  'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r':  'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'v':  'urn:schemas-microsoft-com:vml',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    }

    all_paras = doc.paragraphs
    docpr_id = 900000  # start from high number to avoid conflicts

    # Group OCR results by para_index for efficient processing
    para_ocr_map = {}  # para_index -> list of ocr_results
    for ocr_result in image_ocr_results:
        pi = ocr_result.get("para_index")
        if pi is not None:
            para_ocr_map.setdefault(pi, []).append(ocr_result)

    for para_index, ocr_list in para_ocr_map.items():
        if para_index >= len(all_paras):
            _logger.warning(
                "Image OCR: para_index %s out of range (%d paragraphs)",
                para_index, len(all_paras),
            )
            continue

        para = all_paras[para_index]
        para_element = para._element

        # ── Find all images in this paragraph and their display sizes ──
        # Build a list of (image_width_emu, image_height_emu, base_h_emu, base_v_emu)
        # for each image in the paragraph (in document order).
        images_in_para = []

        # Modern drawings: <w:drawing>
        for drawing in para_element.iter('{%s}drawing' % NSMAP['w']):
            inline_el = drawing.find('{%s}inline' % NSMAP['wp'])
            anchor_el = drawing.find('{%s}anchor' % NSMAP['wp'])
            container = inline_el if inline_el is not None else anchor_el
            if container is None:
                continue

            # Check if this is an actual image (has blip)
            blip = drawing.find('.//{%s}blip' % NSMAP['a'])
            if blip is None:
                continue

            # Get display size from <wp:extent>
            extent = container.find('{%s}extent' % NSMAP['wp'])
            if extent is None:
                continue
            img_cx = int(extent.get('cx', '0'))
            img_cy = int(extent.get('cy', '0'))
            if not img_cx or not img_cy:
                continue

            # Get position for anchor images
            base_h_emu = 0
            base_v_emu = 0
            if anchor_el is not None:
                posH = anchor_el.find('{%s}positionH' % NSMAP['wp'])
                posV = anchor_el.find('{%s}positionV' % NSMAP['wp'])
                if posH is not None:
                    offset_el = posH.find('{%s}posOffset' % NSMAP['wp'])
                    if offset_el is not None and offset_el.text:
                        base_h_emu = int(offset_el.text)
                if posV is not None:
                    offset_el = posV.find('{%s}posOffset' % NSMAP['wp'])
                    if offset_el is not None and offset_el.text:
                        base_v_emu = int(offset_el.text)

            images_in_para.append({
                'cx': img_cx,
                'cy': img_cy,
                'base_h': base_h_emu,
                'base_v': base_v_emu,
                'is_anchor': anchor_el is not None,
            })

        # Legacy VML images
        for imgdata in para_element.iter('{%s}imagedata' % NSMAP['v']):
            parent_shape = imgdata.getparent()
            if parent_shape is None:
                continue
            style_attr = parent_shape.get('style', '')
            w_match = re.search(r'width:\s*([\d.]+)pt', style_attr)
            h_match = re.search(r'height:\s*([\d.]+)pt', style_attr)
            if not w_match or not h_match:
                continue
            # Convert pt to EMU (1 pt = 12700 EMU)
            img_cx = round(float(w_match.group(1)) * 12700)
            img_cy = round(float(h_match.group(1)) * 12700)
            base_h_emu = 0
            base_v_emu = 0
            ml_match = re.search(r'margin-left:\s*([\d.-]+)pt', style_attr)
            mt_match = re.search(r'margin-top:\s*([\d.-]+)pt', style_attr)
            if ml_match:
                base_h_emu = round(float(ml_match.group(1)) * 12700)
            if mt_match:
                base_v_emu = round(float(mt_match.group(1)) * 12700)
            images_in_para.append({
                'cx': img_cx,
                'cy': img_cy,
                'base_h': base_h_emu,
                'base_v': base_v_emu,
                'is_anchor': 'position:absolute' in style_attr,
            })

        if not images_in_para:
            _logger.warning(
                "Image OCR: no images found in paragraph %s XML", para_index,
            )
            continue

        # ── Calculate cumulative horizontal offset for inline images ──
        # Inline images flow left-to-right in the paragraph.
        # Each subsequent inline image's textbox base must be offset by
        # the total width of all preceding inline images.
        cumulative_h = 0
        for img in images_in_para:
            if not img['is_anchor']:
                img['base_h'] = cumulative_h
                cumulative_h += img['cx']

        # ── Create text box overlays for each OCR result ──
        for ocr_result in ocr_list:
            text_blocks = ocr_result.get("text_blocks", [])
            if not text_blocks:
                continue

            img_idx = ocr_result.get("image_index", 0)
            if img_idx >= len(images_in_para):
                img_idx = 0  # fallback to first image
            img_info = images_in_para[img_idx]

            img_cx = img_info['cx']   # image width in EMU
            img_cy = img_info['cy']   # image height in EMU
            base_h = img_info['base_h']
            base_v = img_info['base_v']

            for block in text_blocks:
                translated = block.get("translated", "").strip()
                if not translated:
                    continue

                x_pct = max(0, min(100, float(block.get("x_pct", 0))))
                y_pct = max(0, min(100, float(block.get("y_pct", 0))))
                w_pct = max(1, min(100, float(block.get("w_pct", 10))))
                h_pct = max(1, min(100, float(block.get("h_pct", 5))))

                # Calculate text box position and size in EMU
                tb_left = base_h + round((x_pct / 100.0) * img_cx)
                tb_top = base_v + round((y_pct / 100.0) * img_cy)
                tb_width = round((w_pct / 100.0) * img_cx)
                tb_height = round((h_pct / 100.0) * img_cy)

                # ── Smart font-size based on box area AND text length ──
                # Convert box to points (1 pt = 12700 EMU)
                tb_w_pt = tb_width / 12700.0
                tb_h_pt = tb_height / 12700.0
                char_count = max(1, len(translated))
                # area_font: font that distributes chars evenly across box area
                # (assuming CJK chars ≈ 1em wide)
                area_font = math.sqrt(tb_w_pt * tb_h_pt / char_count)
                # height_font: max font for at least 1 line
                height_font = tb_h_pt * 0.85
                font_size_pt = max(4, min(36, round(min(area_font, height_font))))
                # Scale down font size by 50%
                font_size_pt = max(4, round(font_size_pt * 0.5))
                font_size_hpt = int(font_size_pt * 2)  # half-points

                docpr_id += 1
                try:
                    textbox_run = _make_textbox_xml(
                        translated, tb_left, tb_top, tb_width, tb_height,
                        font_size_hpt, docpr_id,
                    )
                    para_element.append(textbox_run)
                except Exception as e:
                    _logger.warning(
                        "Failed to create OCR textbox for para %s: %s",
                        para_index, e,
                    )
                    continue

    _logger.info(
        "Inserted OCR text boxes for %d image(s)",
        len(image_ocr_results),
    )


def rebuild_docx_from_original(original_content, rebuild_data):
    """Rebuild translated docx by modifying the original document in-place.

    This approach preserves ALL formatting, positioning, textboxes, images,
    tables, and other complex elements. Only text content is replaced.

    Args:
        original_content: Binary content of the original .docx file.
        rebuild_data: dict with keys:
            - paragraphs (list[dict]): Body paragraph + textbox data
            - header_text (str): Translated header text
            - footer_text (str): Translated footer text

    Returns:
        bytes: Binary content of the modified .docx file.
    """
    if Document is None:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    paragraphs_data = rebuild_data.get("paragraphs", [])
    header_text = rebuild_data.get("header_text", "")
    footer_text = rebuild_data.get("footer_text", "")
    image_ocr_results = rebuild_data.get("image_ocr_results", [])

    # Build lookup maps:
    #   para_index → translated_text  for body paragraphs
    #   (para_index, textbox_index) → translated_text  for textboxes
    #   (table_index, row_index) → [cell_texts]  for table rows
    body_translations = {}  # para_index -> translated_text
    textbox_translations = {}  # (para_index, textbox_index) -> translated_text
    table_row_translations = {}  # (table_index, row_index) -> list of cell texts

    def _split_textbox_markers(text):
        """Robustly split text by [TEXTBOX] markers.

        Handles various LLM response formats:
        - '\n[TEXTBOX]\n'  (expected canonical format)
        - '[TEXTBOX]\n'    (LLM drops leading newline)
        - '\n[TEXTBOX]'    (LLM drops trailing newline)
        - '[TEXTBOX]'      (no newlines around marker)

        Returns (para_text, [textbox_texts])
        """
        import re
        # Split by [TEXTBOX] with optional surrounding whitespace/newlines
        parts = re.split(r'\s*\[TEXTBOX\]\s*', text)
        # Filter: first part is paragraph text, rest are textbox texts
        para_text = parts[0].strip() if parts else ""
        tb_texts = [p.strip() for p in parts[1:]] if len(parts) > 1 else []
        return para_text, tb_texts

    for item in paragraphs_data:
        meta = item.get("style_metadata", {})
        translated = item.get("translated_text", "")

        # Split paragraph text and textbox text(s)
        para_text, tb_texts = _split_textbox_markers(translated)

        if meta.get("para_index") is not None:
            pi = meta["para_index"]

            # Strip auto-numbering prefix from translated text
            # (the original doc retains its numbering XML)
            if meta.get("numbering_prefix"):
                para_text = strip_numbering_prefix(
                    para_text, meta["numbering_prefix"]
                )

            body_translations[pi] = para_text

            # Map textbox translations from combined text
            textboxes = meta.get("textboxes", [])
            for tb_idx, tb in enumerate(textboxes):
                tb_translated = tb_texts[tb_idx] if tb_idx < len(tb_texts) else tb.get("full_text", "")
                textbox_translations[(pi, tb_idx)] = tb_translated

        elif meta.get("is_table_row"):
            # Table row: split translated text by [CELL] and map to
            # (table_index, row_index) → list of cell texts
            translated_cells = re.split(r'\s*\[CELL\]\s*', para_text)
            translated_cells = [c.strip() for c in translated_cells]
            tbl_key = (
                meta.get("table_index"),
                meta.get("row_index"),
            )
            table_row_translations[tbl_key] = translated_cells

    # ── Namespace map (must match extraction) ─────────────────────
    NSMAP = {
        'w':   'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r':   'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'v':   'urn:schemas-microsoft-com:vml',
        'mc':  'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    }

    def _is_in_fallback(elem):
        parent = elem.getparent()
        while parent is not None:
            if parent.tag == '{%s}Fallback' % NSMAP['mc']:
                return True
            parent = parent.getparent()
        return False

    def _find_textbox_contents(para_element):
        """Find all unique <w:txbxContent> elements in a paragraph (skip fallback)."""
        result = []
        seen = set()
        for txbx_content in para_element.iter('{%s}txbxContent' % NSMAP['w']):
            if _is_in_fallback(txbx_content):
                continue
            if id(txbx_content) in seen:
                continue
            seen.add(id(txbx_content))
            # Only include if it has text
            texts = []
            for p in txbx_content.findall('{%s}p' % NSMAP['w']):
                for r in p.findall('{%s}r' % NSMAP['w']):
                    for t in r.findall('{%s}t' % NSMAP['w']):
                        texts.append(t.text or "")
            if "".join(texts).strip():
                result.append(txbx_content)
        return result

    def _accept_tracked_changes(container):
        """Keep inserted text and remove deleted text before replacing content."""
        w_ns = NSMAP['w']
        for del_el in list(container.iter('{%s}del' % w_ns)):
            parent = del_el.getparent()
            if parent is not None:
                parent.remove(del_el)
        for ins_el in list(container.iter('{%s}ins' % w_ns)):
            parent = ins_el.getparent()
            if parent is None:
                continue
            idx = parent.index(ins_el)
            for child in list(ins_el):
                parent.insert(idx, child)
                idx += 1
            parent.remove(ins_el)

    def _replace_textbox_text(txbx_content, translated_text):
        """Replace text in a <w:txbxContent> element with translated text.

        Splits translated text by newlines and maps to existing paragraphs.
        Preserves all XML structure and formatting.
        """
        lines = translated_text.split("\n") if translated_text else [""]
        w_paras = txbx_content.findall('{%s}p' % NSMAP['w'])

        for i, w_para in enumerate(w_paras):
            _accept_tracked_changes(w_para)
            new_text = lines[i] if i < len(lines) else ""
            runs = w_para.findall('{%s}r' % NSMAP['w'])
            if runs:
                # Set text on first run's <w:t>, clear rest
                first_t = runs[0].find('{%s}t' % NSMAP['w'])
                if first_t is not None:
                    first_t.text = new_text
                    # Preserve space attribute
                    first_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                for run in runs[1:]:
                    t_elem = run.find('{%s}t' % NSMAP['w'])
                    if t_elem is not None:
                        t_elem.text = ""

    def _replace_para_text(para, translated_text):
        """Replace text in a python-docx Paragraph while preserving formatting.

        Parses <b>/<i>/<u> tags from translated_text and creates separate runs
        for each formatting segment. The first run's base formatting (font,
        color, size) is used as a template for new runs.

        IMPORTANT: We cannot use run.text = xxx because that calls
        clear_content() which removes ALL child elements including
        <w:drawing> (images). Instead we directly manipulate <w:t> elements.
        """
        from lxml import etree
        from copy import deepcopy
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        _accept_tracked_changes(para._element)
        runs = para.runs
        if not runs:
            para.add_run(strip_formatting_tags(translated_text or ""))
            return

        # Check if translated text contains formatting tags
        has_tags = bool(re.search(r'<[biu]>', translated_text or ""))

        if not has_tags:
            # No formatting tags → original behaviour: put all text in first run
            plain = strip_formatting_tags(translated_text or "")
            first_done = False
            for run in runs:
                t_elements = run._element.findall('{%s}t' % w_ns)
                for t_el in t_elements:
                    if not first_done:
                        t_el.text = plain
                        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        first_done = True
                    else:
                        t_el.text = ""
            if not first_done:
                t_el = etree.SubElement(runs[0]._element, '{%s}t' % w_ns)
                t_el.text = plain
                t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            return

        # Has formatting tags → parse into segments and create runs
        segments = _parse_formatting_segments(translated_text)

        # Save the first run's rPr as a formatting template (font, color, size etc.)
        first_rPr = runs[0]._element.find('{%s}rPr' % w_ns)
        template_rPr = deepcopy(first_rPr) if first_rPr is not None else None

        # Remove bold/italic/underline from template (we'll add them per-segment)
        if template_rPr is not None:
            for tag_name in ['b', 'i', 'u', 'bCs', 'iCs']:
                el = template_rPr.find('{%s}%s' % (w_ns, tag_name))
                if el is not None:
                    template_rPr.remove(el)

        # Clear ALL existing <w:t> text in all runs
        for run in runs:
            t_elements = run._element.findall('{%s}t' % w_ns)
            for t_el in t_elements:
                t_el.text = ""

        # Put first segment into first run's first <w:t>
        first_seg = segments[0] if segments else {"text": "", "bold": False, "italic": False, "underline": False}

        first_t = None
        for run in runs:
            t_elements = run._element.findall('{%s}t' % w_ns)
            if t_elements:
                first_t = t_elements[0]
                first_run_elem = run._element
                break

        if first_t is None:
            first_t = etree.SubElement(runs[0]._element, '{%s}t' % w_ns)
            first_run_elem = runs[0]._element

        first_t.text = first_seg["text"]
        first_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        # Update first run's formatting based on segment
        _apply_segment_formatting(first_run_elem, first_seg, w_ns)

        # Create additional runs for remaining segments, insert after the first run element
        insert_after = first_run_elem
        for seg in segments[1:]:
            new_run = etree.SubElement(para._element, '{%s}r' % w_ns)

            # Copy template rPr (font, color, size)
            if template_rPr is not None:
                new_rPr = deepcopy(template_rPr)
                new_run.insert(0, new_rPr)

            # Apply segment-specific formatting
            _apply_segment_formatting(new_run, seg, w_ns)

            # Add <w:t> with text
            new_t = etree.SubElement(new_run, '{%s}t' % w_ns)
            new_t.text = seg["text"]
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

            # Move new_run to right after insert_after
            insert_after.addnext(new_run)
            insert_after = new_run

    def _apply_segment_formatting(run_elem, seg, w_ns):
        """Apply bold/italic/underline to a <w:r> element based on segment flags."""
        from lxml import etree
        rPr = run_elem.find('{%s}rPr' % w_ns)
        if rPr is None:
            rPr = etree.SubElement(run_elem, '{%s}rPr' % w_ns)
            run_elem.insert(0, rPr)

        # Bold
        b_el = rPr.find('{%s}b' % w_ns)
        if seg.get("bold"):
            if b_el is None:
                etree.SubElement(rPr, '{%s}b' % w_ns)
        else:
            if b_el is not None:
                rPr.remove(b_el)

        # Italic
        i_el = rPr.find('{%s}i' % w_ns)
        if seg.get("italic"):
            if i_el is None:
                etree.SubElement(rPr, '{%s}i' % w_ns)
        else:
            if i_el is not None:
                rPr.remove(i_el)

        # Underline
        u_el = rPr.find('{%s}u' % w_ns)
        if seg.get("underline"):
            if u_el is None:
                new_u = etree.SubElement(rPr, '{%s}u' % w_ns)
                new_u.set('{%s}val' % w_ns, 'single')
        else:
            if u_el is not None:
                rPr.remove(u_el)

    # ── Open original document ────────────────────────────────────
    doc = Document(io.BytesIO(original_content))

    # ── Replace body paragraph text ───────────────────────────────
    for para_idx, para in enumerate(doc.paragraphs):
        # Replace main paragraph text
        if para_idx in body_translations:
            translated = body_translations[para_idx]
            if translated is not None:
                _replace_para_text(para, translated)

        # Replace textbox text within this paragraph
        textbox_contents = _find_textbox_contents(para._element)
        for tb_idx, txbx_content in enumerate(textbox_contents):
            key = (para_idx, tb_idx)
            if key in textbox_translations:
                translated = textbox_translations[key]
                if translated is not None:
                    _replace_textbox_text(txbx_content, translated)

    # ── Replace table cell text (row-based with [CELL] separator) ─
    if table_row_translations:
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        def _is_vmerge_cont(tc_elem):
            tcPr = tc_elem.find('{%s}tcPr' % w_ns)
            if tcPr is None:
                return False
            vm = tcPr.find('{%s}vMerge' % w_ns)
            if vm is None:
                return False
            return vm.get('{%s}val' % w_ns) != 'restart'

        for tbl_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                key = (tbl_idx, row_idx)
                if key not in table_row_translations:
                    continue

                translated_cells = table_row_translations[key]

                row_seen = set()
                actual_col = 0
                cell_write_idx = 0  # index into translated_cells
                for cell in row.cells:
                    if id(cell._tc) in row_seen:
                        continue
                    row_seen.add(id(cell._tc))
                    actual_col += 1

                    if _is_vmerge_cont(cell._tc):
                        continue

                    if cell_write_idx >= len(translated_cells):
                        break

                    translated = translated_cells[cell_write_idx]
                    cell_write_idx += 1

                    if translated is None:
                        continue

                    # Put translated text into first paragraph, clear rest
                    cell_paras = cell.paragraphs
                    if cell_paras:
                        _replace_para_text(cell_paras[0], translated)
                        for cp in cell_paras[1:]:
                            _replace_para_text(cp, "")

    # ── Replace header/footer text ────────────────────────────────
    try:
        for section in doc.sections:
            if header_text and section.header:
                for p in section.header.paragraphs:
                    if p.text.strip():
                        _replace_para_text(p, header_text)
                        break
            if footer_text and section.footer:
                for p in section.footer.paragraphs:
                    if p.text.strip():
                        _replace_para_text(p, footer_text)
                        break
    except Exception as e:
        _logger.warning("Failed to replace header/footer text: %s", e)

    # ── Insert OCR text boxes for image translation ───────────────
    if image_ocr_results:
        _insert_image_ocr_textboxes(doc, image_ocr_results)

    # ── Save modified document ────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def rebuild_bilingual_docx_from_original(original_content, bilingual_data):
    """Build a bilingual DOCX by cloning the original document structure.

    The original content is kept in place. For each translated body paragraph,
    a cloned paragraph is inserted immediately after the source paragraph and
    only the clone's text is replaced. For tables, the whole source table is
    cloned once after the source table and translated row/cell text is written
    into the cloned table. This preserves the original template, section setup,
    styles, tables, images, and most document-level formatting.
    """
    if Document is None:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    from copy import deepcopy
    from lxml import etree
    from docx.text.paragraph import Paragraph

    paragraph_translations = bilingual_data.get("paragraphs", {})
    table_translations = bilingual_data.get("tables", {})

    NSMAP = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    }
    w_ns = NSMAP['w']

    def _accept_tracked_changes(container):
        for del_el in list(container.iter('{%s}del' % w_ns)):
            parent = del_el.getparent()
            if parent is not None:
                parent.remove(del_el)
        for ins_el in list(container.iter('{%s}ins' % w_ns)):
            parent = ins_el.getparent()
            if parent is None:
                continue
            idx = parent.index(ins_el)
            for child in list(ins_el):
                parent.insert(idx, child)
                idx += 1
            parent.remove(ins_el)

    def _strip_tags(text):
        return strip_formatting_tags(text or "")

    def _remove_paragraph_numbering(para):
        pPr = para._element.find('{%s}pPr' % w_ns)
        if pPr is None:
            pPr = etree.Element('{%s}pPr' % w_ns)
            para._element.insert(0, pPr)
        numPr = pPr.find('{%s}numPr' % w_ns)
        if numPr is not None:
            pPr.remove(numPr)
        numPr = etree.SubElement(pPr, '{%s}numPr' % w_ns)
        numId = etree.SubElement(numPr, '{%s}numId' % w_ns)
        numId.set('{%s}val' % w_ns, '0')

    def _apply_segment_formatting(run_elem, seg):
        rPr = run_elem.find('{%s}rPr' % w_ns)
        if rPr is None:
            rPr = etree.Element('{%s}rPr' % w_ns)
            run_elem.insert(0, rPr)

        for tag_name, enabled in (("b", seg.get("bold")), ("i", seg.get("italic"))):
            el = rPr.find('{%s}%s' % (w_ns, tag_name))
            if enabled and el is None:
                etree.SubElement(rPr, '{%s}%s' % (w_ns, tag_name))
            elif not enabled and el is not None:
                rPr.remove(el)

        u_el = rPr.find('{%s}u' % w_ns)
        if seg.get("underline"):
            if u_el is None:
                u_el = etree.SubElement(rPr, '{%s}u' % w_ns)
            u_el.set('{%s}val' % w_ns, 'single')
        elif u_el is not None:
            rPr.remove(u_el)

    def _replace_para_text(para, translated_text):
        _accept_tracked_changes(para._element)
        translated_text = translated_text or ""
        runs = para.runs
        if not runs:
            para.add_run(_strip_tags(translated_text))
            return

        has_tags = bool(re.search(r'<[biu]>', translated_text))
        if not has_tags:
            plain = _strip_tags(translated_text)
            first_done = False
            for run in runs:
                for t_el in run._element.findall('{%s}t' % w_ns):
                    if not first_done:
                        t_el.text = plain
                        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        first_done = True
                    else:
                        t_el.text = ""
            if not first_done:
                t_el = etree.SubElement(runs[0]._element, '{%s}t' % w_ns)
                t_el.text = plain
                t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            return

        segments = _parse_formatting_segments(translated_text)
        first_rPr = runs[0]._element.find('{%s}rPr' % w_ns)
        template_rPr = deepcopy(first_rPr) if first_rPr is not None else None
        if template_rPr is not None:
            for tag_name in ['b', 'i', 'u', 'bCs', 'iCs']:
                el = template_rPr.find('{%s}%s' % (w_ns, tag_name))
                if el is not None:
                    template_rPr.remove(el)

        for run in runs:
            for t_el in run._element.findall('{%s}t' % w_ns):
                t_el.text = ""

        first_seg = segments[0] if segments else {"text": "", "bold": False, "italic": False, "underline": False}
        first_t = None
        first_run_elem = runs[0]._element
        for run in runs:
            t_elements = run._element.findall('{%s}t' % w_ns)
            if t_elements:
                first_t = t_elements[0]
                first_run_elem = run._element
                break
        if first_t is None:
            first_t = etree.SubElement(first_run_elem, '{%s}t' % w_ns)

        first_t.text = first_seg["text"]
        first_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        _apply_segment_formatting(first_run_elem, first_seg)

        insert_after = first_run_elem
        for seg in segments[1:]:
            new_run = etree.Element('{%s}r' % w_ns)
            if template_rPr is not None:
                new_run.insert(0, deepcopy(template_rPr))
            _apply_segment_formatting(new_run, seg)
            new_t = etree.SubElement(new_run, '{%s}t' % w_ns)
            new_t.text = seg["text"]
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            insert_after.addnext(new_run)
            insert_after = new_run

    def _cell_texts_for_row(row_data):
        if isinstance(row_data, dict):
            return row_data.get("cells", [])
        return row_data or []

    doc = Document(io.BytesIO(original_content))

    for para_idx, para in list(enumerate(doc.paragraphs)):
        if para_idx not in paragraph_translations:
            continue
        translated = paragraph_translations[para_idx]
        if not _strip_tags(translated).strip():
            continue
        new_p = deepcopy(para._p)
        para._p.addnext(new_p)
        new_para = Paragraph(new_p, para._parent)
        _remove_paragraph_numbering(new_para)
        _replace_para_text(new_para, translated)

    if table_translations:
        for tbl_idx, table in list(enumerate(doc.tables)):
            tbl_key = str(tbl_idx)
            if tbl_key not in table_translations and tbl_idx not in table_translations:
                continue
            rows_data = table_translations.get(tbl_key, table_translations.get(tbl_idx, {}))

            try:
                from docx.table import _Row
                for row_idx, row in reversed(list(enumerate(table.rows))):
                    row_key = str(row_idx)
                    if row_key not in rows_data and row_idx not in rows_data:
                        continue
                    translated_cells = _cell_texts_for_row(rows_data.get(row_key, rows_data.get(row_idx)))
                    if not any(_strip_tags(cell_text).strip() for cell_text in translated_cells):
                        continue

                    new_tr = deepcopy(row._tr)
                    row._tr.addnext(new_tr)
                    translated_row = _Row(new_tr, table)
                    seen_cells = set()
                    write_idx = 0
                    for cell in translated_row.cells:
                        if id(cell._tc) in seen_cells:
                            continue
                        seen_cells.add(id(cell._tc))
                        if write_idx >= len(translated_cells):
                            break
                        text = translated_cells[write_idx]
                        write_idx += 1
                        if cell.paragraphs:
                            _replace_para_text(cell.paragraphs[0], text)
                            for cp in cell.paragraphs[1:]:
                                _replace_para_text(cp, "")
            except Exception:
                _logger.warning("Failed to insert bilingual table rows for table %s", tbl_idx, exc_info=True)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
